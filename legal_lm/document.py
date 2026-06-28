from __future__ import annotations

import re
from pathlib import Path

from legal_lm.schemas import LoadedDocument


class DocumentLoadError(RuntimeError):
    pass


class DocumentLoader:
    supported_extensions = {".txt", ".docx", ".pdf"}

    def load(self, file_path: str | Path) -> LoadedDocument:
        path = Path(file_path)
        if not path.exists():
            raise DocumentLoadError(f"Document not found: {path}")
        suffix = path.suffix.lower()
        if suffix not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise DocumentLoadError(f"Unsupported file type '{suffix}'. Supported: {supported}")

        if suffix == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".docx":
            text = self._load_docx(path)
        else:
            text = self._load_pdf(path)

        text = self._normalize_text(text)
        if not text.strip():
            raise DocumentLoadError(f"No extractable text found in: {path}")

        return LoadedDocument(
            path=str(path),
            file_type=suffix.lstrip("."),
            title=self._extract_title(text, path),
            text=text,
            metadata={"size_bytes": path.stat().st_size},
        )

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentLoadError("python-docx is required to read .docx files.") from exc

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_pdf(self, path: Path) -> str:
        try:
            import PyPDF2
        except ImportError as exc:
            raise DocumentLoadError("PyPDF2 is required to read .pdf files.") from exc

        pages = []
        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            for page in reader.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)

    def _extract_title(self, text: str, path: Path) -> str:
        legal_title_pattern = re.compile(
            r"\b(agreement|contract|amendment|policy|terms|statement|schedule|addendum|notice)\b",
            re.IGNORECASE,
        )
        for raw_line in text.splitlines()[:25]:
            line = raw_line.strip()
            if len(line) < 5:
                continue
            if legal_title_pattern.search(line) or line.isupper():
                return line[:160]
        return path.stem

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
